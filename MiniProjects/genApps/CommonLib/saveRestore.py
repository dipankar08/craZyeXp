###########  Import Export using Excel #####################
try:
    import cPickle as pickle
except:
    import pickle

from openpyxl import Workbook
from openpyxl.compat import range
from openpyxl.cell import get_column_letter

from docx import Document
from docx.shared import Inches


import pdb

class BackUp:
  def __init__(self,ofile_name,otype='pkl',is_encrypt='NO',key=''):
    self.otype = otype
    self.is_encrypt = is_encrypt
    self.key = key
    self.ofile_name =ofile_name
    self.ofull_name=self.ofile_name+'.'+self.otype
  
    #init
    if(self.otype=='pkl'):
      pass
    if(self.otype=='xls'):
      pass
  
  def read(self):
    """ take a file and return list of rows"""
    out = None
    pdb.set_trace()
    if(self.otype=='pkl'):
      with open(self.ofull_name,'rb') as f:
        data = f.read()
        out = pickle.loads(data)
        
    if(self.otype=='xls'):
      from openpyxl import load_workbook
      wb = load_workbook(filename = self.ofull_name)
      sheet_ranges = wb['range names']
      out  =[ [ sheet_ranges['%s%s'%(get_column_letter(col_idx), row)].value  for row in range(1, sheet_ranges.max_column+1) ] for col_idx in range(1, sheet_ranges.max_column+1) ] 

    if(self.otype=='docx'):
      print 'Error: Reading from Docx is not supported'
    if(self.otype=='xls'):
      pass
    if(self.otype=='xls'):
      pass
      
    print 'Success:read'
    return out

  def write(self,data):
    """ Take a list of row and save into a file """
    out = None
    if(self.otype=='pkl'):
      out = pickle.dumps(data)    
      with open(self.ofull_name,'wb') as f:
        f.write(out)
  
    if(self.otype=='xls'):
      wb = Workbook()
      ws = wb.active
      ws.title = "range names"
      for col_idx in range(0, len(data)):
        col = get_column_letter(col_idx+1)
        for row in range(0, len(data[col_idx])):
          ws.cell('%s%s'%(col, row+1)).value = data[col_idx][row]
      ws = wb.create_sheet()
      ws.title = 'Pi'
      ws['F5'] = 3.14
      wb.save(filename = self.ofull_name)
      
    if(self.otype=='docx'):
      document = Document()
      document.add_heading(self.ofull_name, 0)

      table = document.add_table(rows=0, cols=len(data))
      #hdr_cells = table.rows[0].cells
      #hdr_cells[0].text = 'Qty'
      #hdr_cells[1].text = 'Id'
      #hdr_cells[2].text = 'Desc'
      for row in data:
          row_cells = table.add_row().cells
          for col_idx in range(len(row)):
            row_cells[col_idx].text = str(row[col_idx])
      document.add_page_break()
      document.save(self.ofull_name)

    if(self.otype=='xls'):
      pass
    if(self.otype=='xls'):
      pass
 
    print 'Success:Write'




#test code..
#d =[['a','b','c','d'],['a','b','c','d'],['a','b','c','d'],['a','b','c','d']]
#b = BackUp('hello','docx')
#b.write(d)
#print b.read()